from docx import Document
from docx.shared import Inches

import subprocess, sys

guestnames = ['Иван Андреевич', 'Антон Александрович', 'Денис Олегович', 'Василий Сергеевич']
naming = 'автоцентр “Старт”'
address = 'ул. Пушкина 1.'
document = Document()

for x in guestnames:
    document.add_heading('Ремонт Автомобилей', 0)
    p = document.add_paragraph('Уважаемый клиент ')
    p.add_run(x + '! ').bold = True
    document.add_paragraph('')
    p.add_run(' Приглашаем Вас на плановую замену масла ')
    p.add_run('в ' + naming).bold = True
    p.add_run(' по адресу ')
    p.add_run(address).bold = True
    document.add_picture('1.jpeg', width=Inches(5.25))
    document.add_heading('Замена масла бесплатная, фильтр в подарок!', level=1)
    document.add_page_break()

document.save('Sale.docx')

opener = "open" if sys.platform == "darwin" else "xdg-open"
subprocess.call([opener, 'Sale.docx'])
